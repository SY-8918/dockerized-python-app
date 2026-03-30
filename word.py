from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Créer le document Word
doc = Document()

# Titre principal
doc.add_heading("Création d’un tableau de bord des ventes Costco – Version Simple et Humaine", 0)

# Section : Titre de la tâche
doc.add_heading("🧾 Titre de la tâche", level=1)
doc.add_paragraph("👉 Créer un tableau de bord clair avec les ventes des deux sites Costco")

# Section : Qui a besoin de ça ?
doc.add_heading("👤 Qui a besoin de ça ?", level=1)
doc.add_paragraph(
    "Moi, en tant qu’analyste chez Costco, j’en ai marre de passer des heures à ouvrir des fichiers Excel de partout. "
    "On a deux sites (un pour les vêtements et un pour l’électronique) et je dois tout copier-coller, faire des calculs, "
    "convertir des devises… Bref, c’est galère."
)

# Section : Ce qu’on veut vraiment
doc.add_heading("🎯 Ce qu’on veut vraiment", level=1)
doc.add_paragraph(
    "On veut un tableau de bord dans Tableau (le logiciel), qui nous montre en un clin d'œil toutes les ventes des deux sites, "
    "en dollars américains (USD), et avec des données propres. Plus de doublons, plus de cellules vides ou de chiffres bizarres. "
    "Juste les vraies infos, claires, nettes et à jour."
)

# Section : Pourquoi c’est utile ?
doc.add_heading("💡 Pourquoi c’est utile ?", level=1)
doc.add_paragraph(
    "On va gagner un temps fou. Fini le bricolage dans Excel. Les équipes vont pouvoir voir les résultats de vente tout de suite, "
    "sans avoir besoin de demander à quelqu’un de le faire manuellement. Et surtout, les chiffres seront fiables."
)

# Section : Ce qu’on doit faire (en étapes simples)
doc.add_heading("🛠️ Ce qu’on doit faire (en étapes simples)", level=1)
steps = [
    "Aller chercher automatiquement les ventes depuis les deux sites Costco (un pour les fringues, l’autre pour les électroniques).",
    "Nettoyer les données : on enlève les erreurs, les lignes vides, les doublons, etc.",
    "Convertir les montants en USD, parce qu’actuellement certains sont encore en euros.",
    "Créer un joli tableau de bord dans Tableau pour qu’on puisse voir les ventes par produit, par site, etc.",
    "Montrer le tableau à l’équipe commerciale pour voir si ça leur va et faire les petits ajustements s’il faut."
]
for step in steps:
    doc.add_paragraph(f"• {step}", style='List Bullet')

# Section : Quand est-ce qu’on peut dire que c’est bon ?
doc.add_heading("✅ Quand est-ce qu’on peut dire que c’est bon ?", level=1)
success_criteria = [
    "Si les données des deux sites sont bien rassemblées automatiquement,",
    "Si tout est propre et converti en USD,",
    "Si le tableau est clair, à jour, et compréhensible par l’équipe,",
    "Et surtout… si on n’a plus besoin d’ouvrir 5 fichiers Excel chaque matin 😅"
]
for criterion in success_criteria:
    doc.add_paragraph(f"• {criterion}", style='List Bullet')

# Enregistrer le fichier
doc.save("Tableau_de_bord_Costco_simple.docx")
